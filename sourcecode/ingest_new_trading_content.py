#!/usr/bin/env python3
"""
ingest_new_trading_content.py

Ingests the 7 trading/finance files flagged in onedrive_analysis_report.md
into the `trading_lessons` ChromaDB collection.

Companion script to ingest_new_esl_content.py — same structure, same
extraction/chunking logic, different manifest and target collection.

`trading_lessons` currently holds only 21 chunks from a single source
(JETS Options Strangle.pdf), so this meaningfully expands it.

USAGE:
    python3 ingest_new_trading_content.py            # dry run (default)
    python3 ingest_new_trading_content.py --live      # actually writes

Run inside chroma_env, on your own machine (not inside a Cowork VM), so
it can see /Users/gene/Documents/RAG/chroma directly.

Dependencies: chromadb, python-docx, pypdf (same as the ESL script).
"""

import argparse
import re
import sys
from pathlib import Path

from chromadb import PersistentClient
from chromadb.utils import embedding_functions
from docx import Document
from pypdf import PdfReader

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

SOURCE_DIR = Path("/Users/gene/Library/CloudStorage/OneDrive-Personal/Documents")
CHROMA_PATH = "/Users/gene/Documents/RAG/chroma"
COLLECTION_NAME = "trading_lessons"

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150

# ---------------------------------------------------------------------------
# Manifest — the 7 trading/finance files from the report's category (a).
# None of these were close to the existing JETS Options Strangle.pdf source
# (best distance ~0.68, well above the ~0.35 duplicate threshold), so no
# dedup concerns here — all 7 are genuinely new to this collection.
# ---------------------------------------------------------------------------

TRADING_FILES = [
    "Finance English – Set 1.docx",
    "Finance English – Set 3.docx",
    "Introduction to Options – Student Lesson Handout.docx",
    "Investing Learning Experience Step 1.docx",
    "LECCIÓN SOBRE DIVIDENDOS.docx",
    "Options Day 1.docx",
    "Quick Fidelity Stock Guide.docx",
]

MANIFEST = TRADING_FILES

# ---------------------------------------------------------------------------
# Extraction (identical to the ESL script)
# ---------------------------------------------------------------------------

def extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix == ".pdf":
        return extract_pdf_text(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP):
    text = text.strip()
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap
    return chunks


def safe_id_prefix(filename: str) -> str:
    stem = Path(filename).stem
    return re.sub(r"[^A-Za-z0-9_-]", "_", stem)

# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--live",
        action="store_true",
        help="Actually write to ChromaDB. Without this flag, runs as a dry run.",
    )
    args = parser.parse_args()

    print(f"Mode: {'LIVE — will write to ChromaDB' if args.live else 'DRY RUN — no writes'}")
    print(f"Source folder: {SOURCE_DIR}")
    print(f"Chroma path:   {CHROMA_PATH}")
    print(f"Collection:    {COLLECTION_NAME}")
    print(f"Files in manifest: {len(MANIFEST)}\n")

    client = PersistentClient(path=CHROMA_PATH)
    ef = embedding_functions.DefaultEmbeddingFunction()
    collection = client.get_collection(name=COLLECTION_NAME, embedding_function=ef)

    total_chunks = 0
    skipped = []
    errors = []

    for filename in MANIFEST:
        path = SOURCE_DIR / filename
        if not path.exists():
            skipped.append(filename)
            print(f"  [SKIP - not found] {filename}")
            continue

        try:
            text = extract_text(path)
        except Exception as exc:
            errors.append((filename, str(exc)))
            print(f"  [ERROR] {filename}: {exc}")
            continue

        if not text.strip():
            skipped.append(filename)
            print(f"  [SKIP - no extractable text] {filename}")
            continue

        chunks = chunk_text(text)
        prefix = safe_id_prefix(filename)
        ids = [f"{prefix}__{i}" for i in range(len(chunks))]
        metadatas = [
            {"source": filename, "ingested_from": "onedrive_analysis_2026-08-06"}
            for _ in chunks
        ]

        if args.live:
            collection.add(documents=chunks, ids=ids, metadatas=metadatas)

        total_chunks += len(chunks)
        print(f"  [OK] {filename} -> {len(chunks)} chunks")

    print("\n--- Summary ---")
    print(f"Files processed: {len(MANIFEST) - len(skipped) - len(errors)}")
    print(f"Skipped:         {len(skipped)}")
    print(f"Errors:          {len(errors)}")
    print(f"Total chunks:    {total_chunks} ({'written' if args.live else 'would be written'})")

    if skipped:
        print("\nSkipped files (check filenames against actual OneDrive contents):")
        for f in skipped:
            print(f"  - {f}")

    if errors:
        print("\nErrors:")
        for f, e in errors:
            print(f"  - {f}: {e}")

    if not args.live:
        print("\nThis was a dry run. Re-run with --live to actually write to ChromaDB.")


if __name__ == "__main__":
    sys.exit(main())
