#!/usr/bin/env python3
"""
ingest_new_esl_content.py

Ingests the OneDrive files flagged as "(a) relevant, not yet represented"
in onedrive_analysis_report.md into the `esl_lessons` ChromaDB collection.

SCOPE: This script only targets `esl_lessons`. The 7 trading/finance files
from the report (Finance English Set 1/3, Options Day 1, etc.) are NOT
included — those belong in `trading_lessons` instead. Say the word if you
want a second script for that collection.

DEDUPLICATION: The report's "19 CMA/MA Study files" included 5 internal
duplicate pairs (same content saved as both .docx and .pdf). This
manifest lists only ONE file per pair — .docx preferred over .pdf, since
python-docx extraction is generally cleaner than pypdf's. That's a
judgment call; swap the filename in the manifest if you'd rather ingest
the .pdf version for consistency with how other weeks were originally
ingested.

VERIFY BEFORE RUNNING: A few filenames in the "vocabulary/grammar drills"
group (Set 2.docx, Set 5.docx, Set 6.docx) were listed in shorthand in a
comma-separated list in the report. Confirm these are really the exact
on-disk filenames (`ls` the folder or check Finder) before running —
if the real names are longer (e.g. "Clinic and Office English - Set 2.docx"),
fix the manifest below or the script will just skip them with a warning.

USAGE:
    python ingest_new_esl_content.py            # dry run (default) — prints
                                                  # what would happen, writes nothing
    python ingest_new_esl_content.py --live      # actually writes to ChromaDB

Run this inside your chroma_env (the `chroma` alias in your .zshrc), on
your own machine — not inside a Cowork VM. It needs to see your real
ChromaDB files at /Users/gene/Documents/RAG/chroma.

Dependencies (should already be present in chroma_env, since the same
libraries back your existing ingestion scripts):
    pip install chromadb python-docx pypdf
"""

import argparse
import re
import sys
from pathlib import Path

from chromadb import PersistentClient
from chromadb.utils import embedding_functions
from docx import Document
from pypdf import PdfReader

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

SOURCE_DIR = Path("/Users/gene/Library/CloudStorage/OneDrive-Personal/Documents")
CHROMA_PATH = "/Users/gene/Documents/RAG/chroma"
COLLECTION_NAME = "esl_lessons"

# Character-based sliding window. Simple and robust, but can split mid-
# sentence at chunk boundaries — a paragraph-aware chunker would preserve
# semantic units better but produce uneven chunk sizes. Adjust these two
# numbers if you want chunk granularity closer to what ESLPDFIngestion.py
# already uses for the rest of the collection (I don't have that script's
# exact settings, so these are reasonable defaults, not a matched value).
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150

# ---------------------------------------------------------------------------
# Manifest — grouped to match the report's categories, for easy review.
# Files already embedded (report category "b") are deliberately excluded.
# ---------------------------------------------------------------------------

CMA_MA_STUDY_FILES = [
    "Week 1 Day 1 MA Study.docx",       # dup of MA Study Week 1 Day 1.pdf — skipped
    "Week 1 Day 2.docx",
    "Week 1 Day 3 MA Study.docx",
    "Week 1 Day 5 MA Study.docx",
    "Week 2 Day 1 MA English.docx",
    "Week 2 Day 4.docx",
    "Week 2 Day 5.docx",
    "Week 2 Day 6 MA Study.docx",       # dup .pdf skipped
    "Week 3 Day 1 MA Study.docx",       # dup .pdf skipped
    "Week 4 Day 4.docx",                # dup .pdf skipped — fills gap in Week 4
    "Week 6 Day 1.docx",                # dup .pdf skipped — fills gap in Week 6
    "Week 6 Practice.docx",
    "week28.docx",
    "exam-content-outline-effective.pdf",   # only a PDF exists for this one
]

MARISOL_SERIES_FILES = [
    "Advanced verbs ser estar get.docx",
    "CHAPTER 11.docx",
    "Capitulo 12.docx",
    "Capitulo 3 La Revaluación.docx",
    "Captiulo 9.docx",
    "Capítulo 10 El viaje largo.docx",
    "Chapter 14.docx",
    "Chapter 15.docx",
    "Chapter 16  Winter Weather.docx",
    "Chapter 4 La LLamada De La Escuela.docx",
    "Chapter2_Bilingual_Vocab_Pronunciation.docx",
    "Marisol_Chapter5_Bilingual.docx",
    "Marisol_Clinic_Stories_and_Vocab Capitulo 6.docx",
    "Marisol_Friday_Expanded_Lesson Capitulo 8.docx",
    "Marisol_Weekend_Lesson Capitulo 9.docx",
    "The_Retest_Story.docx",
    "The_Retest_Story_Chapter2.docx",
]

VOCAB_DRILL_FILES = [
    "Clinic and Office English – Set 1.docx",
    "Clinic and Office English – Set 2.docx",                        
    "Clinic and Office English – Set 5.docx",                        
    "Clinic and Office English – Set 6.docx",                        
    "Doctors Office Set 3 and 4.docx",
    "Clinic_Story_and_Drills Capitulo 7.docx",
    "English Ser Estar Get.docx",
    "Exercise 12.1 and 12.2.docx",
    "Practice words with W.docx",
    "Reflexive Verbs 12.3.docx",
    "Spanish  English Phrases.docx",
    "Spanish English 3.docx",
    "Speech Practice 2.docx",
    "Speech Practice 3.docx",
    "Test Sentences.docx",
    "Word Pairs.docx",
]

MANIFEST = CMA_MA_STUDY_FILES + MARISOL_SERIES_FILES + VOCAB_DRILL_FILES

# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------

def extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix == ".pdf":
        return extract_pdf_text(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP):
    text = text.strip()
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap
    return chunks


def safe_id_prefix(filename: str) -> str:
    stem = Path(filename).stem
    return re.sub(r"[^A-Za-z0-9_-]", "_", stem)

# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--live",
        action="store_true",
        help="Actually write to ChromaDB. Without this flag, runs as a dry run.",
    )
    args = parser.parse_args()

    print(f"Mode: {'LIVE — will write to ChromaDB' if args.live else 'DRY RUN — no writes'}")
    print(f"Source folder: {SOURCE_DIR}")
    print(f"Chroma path:   {CHROMA_PATH}")
    print(f"Collection:    {COLLECTION_NAME}")
    print(f"Files in manifest: {len(MANIFEST)}\n")

    client = PersistentClient(path=CHROMA_PATH)
    # Same embedding model the report identified esl_lessons as already using.
    ef = embedding_functions.SentenceTransformerEmbeddingFunction(model_name="BAAI/bge-m3")
    collection = client.get_collection(name=COLLECTION_NAME, embedding_function=ef)

    total_chunks = 0
    skipped = []
    errors = []

    for filename in MANIFEST:
        path = SOURCE_DIR / filename
        if not path.exists():
            skipped.append(filename)
            print(f"  [SKIP - not found] {filename}")
            continue

        try:
            text = extract_text(path)
        except Exception as exc:
            errors.append((filename, str(exc)))
            print(f"  [ERROR] {filename}: {exc}")
            continue

        if not text.strip():
            skipped.append(filename)
            print(f"  [SKIP - no extractable text] {filename}")
            continue

        chunks = chunk_text(text)
        prefix = safe_id_prefix(filename)
        ids = [f"{prefix}__{i}" for i in range(len(chunks))]
        metadatas = [
            {"source": filename, "ingested_from": "onedrive_analysis_2026-08-06"}
            for _ in chunks
        ]

        if args.live:
            collection.add(documents=chunks, ids=ids, metadatas=metadatas)

        total_chunks += len(chunks)
        print(f"  [OK] {filename} -> {len(chunks)} chunks")

    print("\n--- Summary ---")
    print(f"Files processed: {len(MANIFEST) - len(skipped) - len(errors)}")
    print(f"Skipped:         {len(skipped)}")
    print(f"Errors:          {len(errors)}")
    print(f"Total chunks:    {total_chunks} ({'written' if args.live else 'would be written'})")

    if skipped:
        print("\nSkipped files (check filenames against actual OneDrive contents):")
        for f in skipped:
            print(f"  - {f}")

    if errors:
        print("\nErrors:")
        for f, e in errors:
            print(f"  - {f}: {e}")

    if not args.live:
        print("\nThis was a dry run. Re-run with --live to actually write to ChromaDB.")


if __name__ == "__main__":
    sys.exit(main())
